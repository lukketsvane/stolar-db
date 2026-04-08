#!/usr/bin/env python3
"""
Replace the Føreord body with Iver's new version.
"""
from __future__ import annotations
import sys
from pathlib import Path
from copy import deepcopy

if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / 'FORMLÆRE.docx'

NEW_FOREORD = [
    "Kvifor har objekt den spesifikke morfologien dei har?",
    "Spørsmålet krev ein dekomposisjon av dei kreftene som verkar. Louis Sullivans diktum om at form følgjer funksjon vert her forkasta som ein logisk blindveg. Jan Michl har demonstrert at dette svaret er ein aktiv distraksjon: form følgjer form. Kvar ny aktualisering er ein modifikasjon av ein eksisterande posisjon i formrommet, utført av agentar med avgrensa kognitiv lyskjegle. Michl rydda det intellektuelle landskapet, men han formulerte berre det negative resultatet: han viste kva form ikkje følgjer.",
    "George Kubler (1962) nådde nærare ved å skildre formelle sekvensar og «prime objects». Han såg formhistoria som eit problem om posisjon og temporalitet, men han mangla eit formelt apparate for å identifisere kreftene som driv sekvensane. Han kartla rørsla utan å identifisere gradientane.",
    "Denne traktaten er eit forsøk på å etablere det Kubler og Michl let stå ope: eit substrat-uavhengig rammeverk for korleis form emergerer. Ambisjonen er ikkje å forklare einskildobjektet, men den generelle prosessen som genererer alle punkt i formrommet; frå biologiske organ til generative algoritmar.",
    "Form er ein posisjon i eit tilpassingslandskap, generert av aggregerte seleksjonstrykk og navigert av agentar på fleire hierarkiske skalaer.",
    "Systemavgrensingar. Traktaten skildrar krefter og posisjonar, ikkje verdiar. Ho kan forklare avstanden mellom ein Wegner-stol og ein Monobloc i formrommet, men ho kan ikkje felle estetiske domar. Estetikk er transcendent til formsystemet, på same vis som etikk er transcendent til logikken (Wittgenstein, 1921). Ikkje fordi det er uviktig, men fordi det ikkje let seg fange i proposisjonar.",
    "Traktaten seier heller ikkje noko om intensjon. Formgjevaren si oppleving av å skape er utilgjengeleg for modellen. Agenten navigerer; modellen registrerer berre den resulterande posisjonen i landskapet. Innvendig liv og ytre spor er to ulike skildringsnivå.",
    "Med tilvising til Wittgenstein: den som forstår desse proposisjonane, må erkjenne dei som ein stige. Proposisjonane skal ikkje erstatte handverket, men gjere den intuitive navigasjonen eksplisitt. Rammeverket skal vise agenten kva han allereie gjer.",
    "Dersom denne skildringa er adekvat, fangar ho eigenskapar ved formverda som var sanne før dei vart formulerte, og som vil vere sanne etter at orda er gløymde.",
    "Oslo, 2026",
]


def main() -> int:
    d = Document(str(DOCX))

    # Find Føreord heading and the next H1 (Ordliste)
    foreord_idx = None
    ordliste_idx = None
    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if t == 'Føreord' and foreord_idx is None:
            foreord_idx = i
        elif t == 'Ordliste' and foreord_idx is not None:
            ordliste_idx = i
            break
    if foreord_idx is None or ordliste_idx is None:
        print('ERROR: could not find Føreord/Ordliste boundary', file=sys.stderr)
        return 1

    # Delete everything between (exclusive)
    paras = list(d.paragraphs)
    template = paras[foreord_idx + 1]  # use first body paragraph as style template
    template_el = deepcopy(template._p)
    # strip runs from template
    for child in list(template_el):
        if child.tag == qn('w:r'):
            template_el.remove(child)

    foreord_el = paras[foreord_idx]._p
    parent = foreord_el.getparent()

    # Delete body (between Føreord heading and Ordliste heading, exclusive)
    to_delete = []
    for i in range(foreord_idx + 1, ordliste_idx):
        to_delete.append(paras[i]._p)
    print(f'Deleting {len(to_delete)} old Føreord body paragraphs')
    for el in to_delete:
        el.getparent().remove(el)

    # Now insert new paragraphs after Føreord heading
    # Re-find foreord position (indices shifted)
    foreord_pos = list(parent).index(foreord_el)
    for offset, text in enumerate(NEW_FOREORD, start=1):
        new_el = deepcopy(template_el)
        # Add a single run with the text
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        new_el.append(r)
        parent.insert(foreord_pos + offset, new_el)

    d.save(str(DOCX))
    print(f'Inserted {len(NEW_FOREORD)} new Føreord paragraphs')
    return 0


if __name__ == '__main__':
    sys.exit(main())
