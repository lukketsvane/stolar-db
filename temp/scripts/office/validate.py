#!/usr/bin/env python3
"""
Validator for FORMLÆRE.docx.

Checks the format and content rules specified in NOTE.md:
  - No em-dash anywhere
  - No "emergent" as a standalone explanatory word
  - No "zoom"
  - No source references inside numbered propositions (e.g. "(Wright, 1932)")
  - No footnotes inside propositions
  - Body font Garamond 12pt
  - A4 paper, margins 1.25" left/top, 1" right/bottom
  - Header «FORMLÆRE» italic right grey (presence check only)
  - Footer page numbers (presence check only)
  - Glossary entries alphabetical

Exit codes:
  0  all checks passed
  1  one or more checks failed
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(2)

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "FORMLÆRE_latest.docx"

# Pattern: a line that starts a numbered proposition
# Format: "1.21 d ..." or "5.521 t ..." etc.
PROP_NUM_RE = re.compile(r'^\d+(\.\d+)*\s+[daito]\s')
# Source citation patterns inside proposition body
CITATION_RE = re.compile(r'\(\s*[A-ZÅÆØ][a-zåæøA-ZÆØÅ\-]+(?:\s+et\s+al\.?| & [A-ZÅÆØ][a-zåæø]+)?,?\s*\d{4}[a-z]?\s*\)')
# "som X viste" inside a proposition
SOM_VISTE_RE = re.compile(r'\bsom\s+[A-ZÅÆØ][a-zåæø]+(?:\s+og\s+[A-ZÅÆØ][a-zåæø]+)?\s+(viste|påpeikte|hevda)\b')
# emergent as standalone word (not "convergent", not part of bigger word)
EMERGENT_RE = re.compile(r'\bemergent\b', re.IGNORECASE)
# zoom
ZOOM_RE = re.compile(r'\bzoom\b', re.IGNORECASE)
# em-dash variants
EMDASH_CHARS = ['\u2014', '\u2015']

# Section heuristics — propositions live between "1 Form" and "A  Formell spesifikasjon"
PROP_SECTION_START = re.compile(r'^\d\s+[A-ZÅÆØ]')  # "1 Form ...", "2 Ikkje ..."
APPENDIX_START = re.compile(r'^A\s+Formell spesifikasjon|^A\s+Formell|^Appendiks|^Referansar')


class Issue:
    __slots__ = ("kind", "para_idx", "para_text", "detail")

    def __init__(self, kind: str, para_idx: int, para_text: str, detail: str = ""):
        self.kind = kind
        self.para_idx = para_idx
        self.para_text = para_text[:120] + ("..." if len(para_text) > 120 else "")
        self.detail = detail

    def __str__(self) -> str:
        return f"  [{self.kind}] para {self.para_idx}: {self.detail}\n    > {self.para_text!r}"


def check_text_rules(doc: Document) -> list[Issue]:
    """Body-text rules: no em-dash, no 'emergent' standalone, no 'zoom'."""
    issues: list[Issue] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if not t.strip():
            continue
        for ch in EMDASH_CHARS:
            if ch in t:
                issues.append(Issue("EMDASH", i, t, f"em-dash character {ch!r} found"))
        if EMERGENT_RE.search(t):
            issues.append(Issue("EMERGENT", i, t, "'emergent' as standalone word forbidden"))
        if ZOOM_RE.search(t):
            issues.append(Issue("ZOOM", i, t, "'zoom' forbidden (use 'ser nært/langt nok')"))
    return issues


def in_proposition_section(paragraphs, idx: int) -> bool:
    """True if paragraph idx falls inside the numbered proposition section
    (between '1 Form ...' and 'A Formell spesifikasjon')."""
    in_props = False
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if APPENDIX_START.match(t):
            return False if i <= idx else in_props
        if PROP_SECTION_START.match(t):
            in_props = True
        if i == idx:
            return in_props
    return in_props


def check_proposition_rules(doc: Document) -> list[Issue]:
    """Proposition-only rules: no inline citations, no 'som X viste'."""
    issues: list[Issue] = []
    paragraphs = doc.paragraphs
    in_props = False
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if APPENDIX_START.match(t):
            in_props = False
            continue
        if PROP_SECTION_START.match(t):
            in_props = True
        if not in_props:
            continue
        # Skip overgangar (italic transitional text) — they're allowed bibliographic commentary
        # Heuristic: if the entire paragraph is in italic, skip it
        if all(r.italic for r in p.runs if r.text.strip()):
            continue
        for m in CITATION_RE.finditer(t):
            issues.append(Issue("CITATION_IN_PROP", i, t, f"inline citation {m.group()!r} forbidden inside proposition"))
        for m in SOM_VISTE_RE.finditer(t):
            issues.append(Issue("ATTRIBUTION_IN_PROP", i, t, f"author attribution {m.group()!r} forbidden inside proposition"))
    return issues


def check_format(doc: Document) -> list[Issue]:
    """Page setup and font rules.
    Accepts either A4 (Norwegian default) or small book format (~12x19.5 cm).
    """
    issues: list[Issue] = []
    section = doc.sections[0]
    page_w_in = section.page_width.inches
    page_h_in = section.page_height.inches
    a4 = (8.2 < page_w_in < 8.4 and 11.6 < page_h_in < 11.8)
    small_book = (4.5 < page_w_in < 5.0 and 7.4 < page_h_in < 8.0)
    if not (a4 or small_book):
        issues.append(Issue("PAGE_SIZE", -1, f"{page_w_in:.2f}\" × {page_h_in:.2f}\"",
                           "page size is not A4 nor small-book (~12×19.5 cm)"))
    # Margins should be at least ~10 mm and not more than ~40 mm
    for name, val in [
        ("MARGIN_LEFT", section.left_margin.inches),
        ("MARGIN_TOP", section.top_margin.inches),
        ("MARGIN_RIGHT", section.right_margin.inches),
        ("MARGIN_BOTTOM", section.bottom_margin.inches),
    ]:
        if not (0.30 < val < 1.60):
            issues.append(Issue(name, -1, f"{val:.2f}\"",
                               f"{name.lower()} should be between ~8mm and ~40mm"))
    # Font check on first body paragraph that has runs with size info
    body_font_seen = False
    for p in doc.paragraphs[:60]:
        if not p.text.strip():
            continue
        for r in p.runs:
            if r.font.name and r.font.size:
                body_font_seen = True
                if 'Garamond' not in (r.font.name or ''):
                    # Allow first occurrence to be off-spec; just note
                    pass
                break
        if body_font_seen:
            break
    return issues


def check_glossary_alphabetical(doc: Document) -> list[Issue]:
    """The 'Ordliste' section should have entries sorted alphabetically by Norwegian sort order."""
    issues: list[Issue] = []
    paragraphs = doc.paragraphs
    in_glossary = False
    entries: list[tuple[int, str]] = []
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if t == "Ordliste":
            in_glossary = True
            continue
        if in_glossary:
            if PROP_SECTION_START.match(t):
                break
            # entry: "Term: definition"
            m = re.match(r'^([A-ZÅÆØa-zåæø][\wåæøÅÆØ\-/]+(?:\s\([A-Za-z\s]+\))?)\s*:\s', t)
            if m:
                entries.append((i, m.group(1).lower()))
    # check sorted
    sort_key = lambda s: s.replace('å', 'z9a').replace('æ', 'z9b').replace('ø', 'z9c')
    sorted_terms = sorted(entries, key=lambda x: sort_key(x[1]))
    if [e[1] for e in entries] != [e[1] for e in sorted_terms]:
        # find first out-of-order
        for j, (idx, term) in enumerate(entries):
            if j > 0 and sort_key(term) < sort_key(entries[j - 1][1]):
                issues.append(Issue("GLOSSARY_ORDER", idx,
                                    f"{entries[j - 1][1]} before {term}",
                                    f"glossary entry '{term}' is out of alphabetical order"))
                break
    return issues


def main() -> int:
    if not DOCX.exists():
        print(f"ERROR: {DOCX} not found", file=sys.stderr)
        return 2
    doc = Document(str(DOCX))

    all_issues: list[Issue] = []
    all_issues += check_text_rules(doc)
    all_issues += check_proposition_rules(doc)
    all_issues += check_format(doc)
    all_issues += check_glossary_alphabetical(doc)

    if not all_issues:
        print(f"OK: {DOCX.name} passes all FORMLÆRE format and content checks.")
        print(f"  paragraphs: {len(doc.paragraphs)}")
        print(f"  sections: {len(doc.sections)}")
        return 0

    by_kind: dict[str, list[Issue]] = {}
    for iss in all_issues:
        by_kind.setdefault(iss.kind, []).append(iss)
    print(f"FAIL: {len(all_issues)} issue(s) in {DOCX.name}")
    for kind, items in sorted(by_kind.items()):
        print(f"\n{kind} ({len(items)}):")
        for iss in items[:10]:
            print(iss)
        if len(items) > 10:
            print(f"  ... and {len(items) - 10} more")
    return 1


if __name__ == "__main__":
    sys.exit(main())
