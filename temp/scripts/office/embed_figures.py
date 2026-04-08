#!/usr/bin/env python3
"""
Embed strong-finding figures inline in FORMLÆRE.docx.

For each (anchor-prefix, figure-path, caption) entry below, find the
matching paragraph and insert image + italic caption AFTER it. Idempotent:
if a caption with the same text already exists, the entry is skipped.

Width: 14 cm to fit A4 with 1.25"/1.0" margins (page text width ≈ 14.7 cm).
"""
from __future__ import annotations
import sys
from pathlib import Path

if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / 'FORMLÆRE_latest.docx'
FIG = ROOT / 'analysis' / 'figures'

# Each entry: (anchor prefix, figure file, italic caption)
# Anchors are matched against the start of paragraph text.
EMBED = [
    (
        '1.4o',
        'fig-1.4-nn-cv.png',
        'Nærmaste-nabo-fordelinga er klumpa, ikkje uniform: CV = 5,4, 15× over Poisson-nullen.',
    ),
    (
        '2.62o',
        'fig-2.4-proxy.png',
        'Stilperiode slår materiale på alle fire dimensjonar: eit samansett trykk er sterkare enn delane.',
    ),
    (
        '3.3d',
        'fig-3.3-channeling.png',
        'Kanaliseringshierarkiet i mesh-rommet: sphericity er sterkt kanalisert, råvolumet fritt, ein 128× spreiing.',
    ),
    (
        '3.4d',
        'fig-3.4-silhouette.png',
        'Stilar er gradientar i mesh-rommet, ikkje topologiske klynger (silhouette = −0,34).',
    ),
    (
        '4.4o',
        'fig-4.4-hull.png',
        'Det kumulative formromsvolumet veks monotont gjennom 700 år (totalvekst 107×).',
    ),
    (
        '4.5i',
        'fig-4.5-mahogni.png',
        'Norsk mahogni-kollapsen 1825-1849: eitt seleksjonstrykk overkøyrer alle andre.',
    ),
]

# Etterord: place the falsification figure near the falsification paragraph
ETTERORD_EMBED = [
    (
        'attekstenkanfellast',
        'fig-falsification-4.1.png',
        'Wasserstein-distanse mellom suksessive 50-årsperiodar: landskapet endrar seg overalt, postulat 4.1 held.',
    ),
]


def find_paragraph_by_prefix(doc, prefix: str):
    import re
    # Normalize anchor for matching: '1.4 o' -> '1.4o'
    norm = re.sub(r'\s+', '', prefix).strip().lower()
    for i, p in enumerate(doc.paragraphs):
        # Normalize paragraph text for matching prefix
        p_text_norm = re.sub(r'\s+', '', p.text[:100]).strip().lower()
        if p_text_norm.startswith(norm):
            return i, p
    return None, None


def insert_image_and_caption_after(doc: 'Document', anchor_p, image_path: Path, caption: str) -> bool:
    """Insert an image paragraph and a caption paragraph after anchor_p.
    Returns True if inserted, False if already present (idempotent)."""
    # Idempotent check: scan up to 4 paragraphs after anchor for the same caption
    parent = anchor_p._p.getparent()
    anchor_idx = list(parent).index(anchor_p._p)
    for k in range(1, 5):
        if anchor_idx + k >= len(parent):
            break
        nxt = parent[anchor_idx + k]
        if nxt.tag != qn('w:p'):
            continue
        # extract text quickly
        nxt_text = ''.join(t.text or '' for t in nxt.iter(qn('w:t')))
        if caption in nxt_text:
            return False  # already present

    # Build template paragraph from anchor (preserves style)
    img_p_el = deepcopy(anchor_p._p)
    for child in list(img_p_el):
        if child.tag == qn('w:r') or child.tag == qn('w:fldSimple'):
            img_p_el.remove(child)
    cap_p_el = deepcopy(img_p_el)

    parent.insert(anchor_idx + 1, img_p_el)
    parent.insert(anchor_idx + 2, cap_p_el)

    # Wrap as Paragraph for python-docx convenience
    from docx.text.paragraph import Paragraph
    img_para = Paragraph(img_p_el, anchor_p._parent)
    cap_para = Paragraph(cap_p_el, anchor_p._parent)

    # Image: add as run with picture
    img_run = img_para.add_run()
    img_run.add_picture(str(image_path), width=Cm(8.5))
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.left_indent = None
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(2)

    # Caption: italic, smaller, centered
    cap_run = cap_para.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.name = 'Garamond'
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.left_indent = None
    cap_para.paragraph_format.space_before = Pt(0)
    cap_para.paragraph_format.space_after = Pt(8)
    return True


def main() -> int:
    d = Document(str(DOCX))
    inserted = 0
    skipped = 0
    for prefix, fname, caption in EMBED + ETTERORD_EMBED:
        idx, p = find_paragraph_by_prefix(d, prefix)
        if idx is None:
            print(f'  MISS: anchor not found: {prefix[:60]!r}')
            continue
        path = FIG / fname
        if not path.exists():
            print(f'  MISS: figure file not found: {path}')
            continue
        if insert_image_and_caption_after(d, p, path, caption):
            inserted += 1
            print(f'  + {fname} after {prefix[:50]!r}')
        else:
            skipped += 1
            print(f'  = {fname} already present (skipped)')
    d.save(str(DOCX))
    print(f'\n{inserted} figures embedded, {skipped} already present')
    return 0


if __name__ == '__main__':
    sys.exit(main())
