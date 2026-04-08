#!/usr/bin/env python3
"""
Embed figures inline inside A.6 sub-sections of FORMLÆRE.docx.

For each A.6.x heading we walk forward to find the body paragraph
(the one after the heading) and insert the figure immediately after
it. The italic caption that already follows is left in place.

Idempotent: if an image already exists between the body and the
caption, the entry is skipped.
"""
from __future__ import annotations
import sys
from pathlib import Path
from copy import deepcopy

if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / 'FORMLÆRE.docx'
FIG = ROOT / 'analysis' / 'figures'

# Map A.6.x heading prefix -> figure file
EMBED = [
    ('A.6.1 Formrommet er ikkje uniformt', 'fig-1.4-morphospace.png'),
    ('A.6.2 Stilperiode som samlevariabel', 'fig-2.4-prediktor.png'),
    ('A.6.3 Kanaliseringshierarki', 'fig-3.3-channeling-v2.png'),
    ('A.6.4 Stilar er gradientar', 'fig-3.4-silhouette.png'),
    ('A.6.5 Kumulativ ekspansjon', 'I-4_morphospace_ekspansjon.png'),
    ('A.6.6 Mahogni-kollapsen', 'fig-4.5-mahogni.png'),
    ('A.6.7 Direkte falsifisering', 'fig-falsification-4.1.png'),
]


def find_paragraph_starting_with(doc, prefix: str):
    norm = prefix.strip()
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(norm):
            return i, p
    return None, None


def has_image_after(parent, idx: int, max_lookahead: int = 4) -> bool:
    """Check if any of the next few paragraphs already contains a drawing."""
    for k in range(1, max_lookahead + 1):
        if idx + k >= len(parent):
            return False
        el = parent[idx + k]
        if el.tag != qn('w:p'):
            continue
        if el.find('.//' + qn('w:drawing')) is not None:
            return True
    return False


def main() -> int:
    d = Document(str(DOCX))
    inserted = 0
    skipped = 0

    for prefix, fname in EMBED:
        idx, heading_p = find_paragraph_starting_with(d, prefix)
        if idx is None:
            print(f'  MISS: heading not found: {prefix[:50]!r}')
            continue
        # Body paragraph is heading + 1
        if idx + 1 >= len(d.paragraphs):
            print(f'  MISS: no body after {prefix[:50]!r}')
            continue
        body_p = d.paragraphs[idx + 1]
        body_el = body_p._p
        parent = body_el.getparent()
        body_pos = list(parent).index(body_el)

        if has_image_after(parent, body_pos):
            skipped += 1
            print(f'  = {fname} already present after {prefix[:50]!r}')
            continue

        path = FIG / fname
        if not path.exists():
            print(f'  MISS: figure file not found: {path}')
            continue

        # Build a new paragraph that inherits the body's style, then add a centred image run
        new_p_el = deepcopy(body_el)
        for child in list(new_p_el):
            if child.tag == qn('w:r') or child.tag == qn('w:fldSimple'):
                new_p_el.remove(child)
        parent.insert(body_pos + 1, new_p_el)

        # Wrap and add picture
        d2 = Document(str(DOCX))  # we'll save and reopen below; for now use python-docx wrapper
        # Use the parent doc to wrap properly
        new_para = Paragraph(new_p_el, body_p._parent)
        run = new_para.add_run()
        run.add_picture(str(path), width=Cm(14))
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_para.paragraph_format.left_indent = None
        new_para.paragraph_format.space_before = Pt(8)
        new_para.paragraph_format.space_after = Pt(2)

        inserted += 1
        print(f'  + {fname} after {prefix[:50]!r}')

    d.save(str(DOCX))
    print(f'\n{inserted} figures embedded, {skipped} already present')
    return 0


if __name__ == '__main__':
    sys.exit(main())
