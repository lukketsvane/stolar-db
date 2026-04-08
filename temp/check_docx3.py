from docx import Document
doc = Document(r'C:\Users\Shadow\Documents\GitHub\stolar-db\FORMLÆRE.docx')
# Show wider context around 316
for k in range(305, 325):
    if k < len(doc.paragraphs):
        style = doc.paragraphs[k].style.name if doc.paragraphs[k].style else ''
        print(f'[{k}] {style}: {doc.paragraphs[k].text[:140]!r}')
