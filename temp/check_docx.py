from docx import Document
doc = Document(r'C:\Users\Shadow\Documents\GitHub\stolar-db\FORMLÆRE.docx')
falsif_count = 0
prop_a_count = 0
import re
PROP_RE = re.compile(r'^(\d+(?:\.\d+)*)\s*([daiot])[\s\t]+(.*)$')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('Falsifiseringsvilkår'):
        falsif_count += 1
        # Show the previous non-empty para
        for j in range(i-1, -1, -1):
            pj = doc.paragraphs[j].text.strip()
            if pj:
                m = PROP_RE.match(pj)
                if m:
                    print(f'  fals after prop {m.group(1)}{m.group(2)}: {t[:80]}')
                else:
                    print(f'  fals after non-prop: prev={pj[:40]} | {t[:60]}')
                break
    if PROP_RE.match(t) and PROP_RE.match(t).group(2) == 'a':
        prop_a_count += 1
print(f'\ntotal Falsifiseringsvilkår paragraphs: {falsif_count}')
print(f'total status-a propositions: {prop_a_count}')
# Look for any remaining Falsifiseringsvilkår mention
for p in doc.paragraphs:
    if 'Falsifiseringsvilkår' in p.text and not p.text.strip().startswith('Falsifiseringsvilkår'):
        print(f'  embedded mention: {p.text[:100]}')
