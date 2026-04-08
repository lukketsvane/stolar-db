from docx import Document
doc = Document(r'C:\Users\Shadow\Documents\GitHub\stolar-db\FORMLÆRE.docx')
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if 'Falsifiseringsvilkår' in t:
        # context: prev 2 + this + next 2
        for k in range(max(0, i-3), min(len(doc.paragraphs), i+4)):
            marker = '>>>' if k == i else '   '
            print(f'{marker} [{k}] {doc.paragraphs[k].text[:120]!r}')
        print('---')
