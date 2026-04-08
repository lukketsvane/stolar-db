import sys, re
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
doc = Document(r'C:\Users\Shadow\Documents\GitHub\stolar-db\FORMLÆRE.docx')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '4 Landskapet' in t or t in {'4 Landskapet er dynamisk.', '4'}:
        style = p.style.name if p.style else '?'
        for k in range(max(0,i-2), min(len(doc.paragraphs), i+5)):
            marker = '>>>' if k == i else '   '
            pt = doc.paragraphs[k].text.strip()
            ps = doc.paragraphs[k].style.name if doc.paragraphs[k].style else '?'
            print(f'{marker}[{k}] {ps}: {pt[:100]!r}')
        print('---')
