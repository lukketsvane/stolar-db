#!/usr/bin/env python3
"""
Surgical-edit utility for FORMLÆRE.docx.

Applies a JSON edit-script of paragraph-level operations:

  [
    {"match": "5.521 t Mengda", "replace": "5.521 t New text..."},
    {"match": "Stale paragraph prefix", "delete": true},
    {"match": "5.6 o Navigasjonskompetanse",
     "insert_after": ["5.61 t New proposition...", "5.62 o Another..."]}
  ]

Replace: finds the first paragraph whose text starts with `match` and
replaces its full text with `replace`. Run formatting is rebuilt as a
single run inheriting the paragraph default.

Delete: finds the first paragraph whose text starts with `match` and
removes the paragraph element from the document tree (no blank line
left behind).

Insert_after: inserts one or more new paragraphs immediately after the
matched paragraph. The new paragraphs inherit the matched paragraph's
style. Order is preserved.

Usage:
  python scripts/office/edit.py --apply edits.json
  python scripts/office/edit.py --dry-run edits.json
  python scripts/office/edit.py --apply edits.json --smart-prop

Exit codes:
  0  success (or dry-run no errors)
  1  one or more match strings not found
  2  invocation/parse error
"""
from __future__ import annotations
import argparse
import json
import re
import sys
from pathlib import Path
from copy import deepcopy

# Force stdout/stderr to UTF-8 on Windows so Unicode (∧, ≤, ø, å, æ) prints cleanly.
if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    print("ERROR: python-docx not installed.", file=sys.stderr)
    sys.exit(2)

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "FORMLÆRE.docx"

# "5.521 t Mengda..." or "T6 (Agent..." or similar.
PROP_LEAD = re.compile(r'^([A-Z]?\d+(?:\.\d+)*)\s+([daiot])\s+(.*)$', re.DOTALL)
THEOREM_LEAD = re.compile(r'^(T\d+|D\d+|A\d+)\s+', re.DOTALL)


def find_paragraph_by_prefix(doc, prefix: str):
    """Return (idx, paragraph) for the first paragraph whose stripped text starts with prefix."""
    norm_prefix = prefix.strip()
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(norm_prefix):
            return i, p
    return None, None


def replace_paragraph_text(p, new_text: str, smart_prop: bool = False) -> None:
    """Replace the entire text of paragraph p with new_text.
    Run formatting is rebuilt: clears existing runs and inserts new ones.
    If smart_prop=True and new_text matches the proposition format
    "N.N x Body", emit:
      run1: "N.N "  bold
      run2: "x"     bold superscript
      run3: " Body" normal
    Otherwise the whole text becomes one normal run.
    """
    # Remember the paragraph style; remove all runs and replace
    # docx Paragraph object: clear via element manipulation
    para_el = p._p
    # Remove existing run elements but keep paragraph properties (pPr)
    for child in list(para_el):
        if child.tag == qn('w:r'):
            para_el.remove(child)

    if smart_prop:
        m = PROP_LEAD.match(new_text)
        if m:
            num, status, body = m.groups()
            r1 = p.add_run(f"{num} ")
            r1.bold = True
            r2 = p.add_run(status)
            r2.bold = True
            r2.font.superscript = True
            r3 = p.add_run(f" {body}")
            return
        m2 = THEOREM_LEAD.match(new_text)
        if m2:
            label = m2.group(1)
            rest = new_text[len(label):]
            r1 = p.add_run(label)
            r1.bold = True
            r2 = p.add_run(rest)
            return
    # Default: single run
    p.add_run(new_text)


def delete_paragraph(p) -> None:
    """Remove the paragraph element from its parent in the XML tree."""
    el = p._p
    el.getparent().remove(el)


def insert_paragraphs_after(p, texts: list[str], smart_prop: bool = False) -> None:
    """Insert new paragraphs after p. Each new paragraph copies the style
    of p (so the new propositions inherit Garamond/normal style)."""
    from copy import deepcopy
    src = p._p
    parent = src.getparent()
    src_idx = list(parent).index(src)
    # Build template paragraph element from p (with formatting but no runs)
    template = deepcopy(src)
    # Strip runs from template, keep paragraph properties (pPr)
    for child in list(template):
        if child.tag == qn('w:r'):
            template.remove(child)
    # Insert new paragraphs in reverse order so each ends up immediately after p
    for offset, text in enumerate(texts, start=1):
        new_el = deepcopy(template)
        parent.insert(src_idx + offset, new_el)
    # Now fill in the text via python-docx wrapper. We need to reconstruct
    # the Paragraph wrappers; easiest is to walk the doc again.
    # Instead, build runs directly from the text on each new element.
    from docx.oxml.ns import qn as _qn
    # Re-fetch the inserted elements (they're at src_idx+1 .. src_idx+len(texts))
    inserted_elements = list(parent)[src_idx + 1: src_idx + 1 + len(texts)]
    for el, text in zip(inserted_elements, texts):
        _add_text_run(el, text, smart_prop=smart_prop)


def _add_text_run(p_el, text: str, smart_prop: bool = False) -> None:
    """Append a w:r run with text to a w:p element. Optionally apply
    smart proposition formatting (bold number + superscript status letter)."""
    from lxml import etree as _et
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    nsmap = {'w': W}

    def make_run(txt: str, bold: bool = False, superscript: bool = False) -> _et._Element:
        r = _et.SubElement(p_el, qn('w:r'))
        if bold or superscript:
            rpr = _et.SubElement(r, qn('w:rPr'))
            if bold:
                _et.SubElement(rpr, qn('w:b'))
            if superscript:
                vert = _et.SubElement(rpr, qn('w:vertAlign'))
                vert.set(qn('w:val'), 'superscript')
        t = _et.SubElement(r, qn('w:t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = txt
        return r

    if smart_prop:
        m = PROP_LEAD.match(text)
        if m:
            num, status, body = m.groups()
            make_run(f"{num} ", bold=True)
            make_run(status, bold=True, superscript=True)
            make_run(f" {body}")
            return
    make_run(text)


def apply_edits(doc, edits: list[dict], smart_prop: bool = False, dry_run: bool = False) -> tuple[int, int]:
    """Apply each edit. Returns (applied, missed).
    Edit ops: `replace` (default), `delete`, `insert_after`.
    Note: deletes/inserts shift later paragraph indices, so we re-resolve
    each edit by prefix match against the current document state."""
    applied = 0
    missed = 0
    for n, e in enumerate(edits, start=1):
        match = e.get("match", "")
        do_delete = bool(e.get("delete", False))
        insert_after = e.get("insert_after")
        replace = e.get("replace", "")
        if not match:
            print(f"[{n}] SKIP: edit missing 'match'", file=sys.stderr)
            missed += 1
            continue
        if not do_delete and not replace and not insert_after:
            print(f"[{n}] SKIP: edit missing 'replace' / 'delete' / 'insert_after'", file=sys.stderr)
            missed += 1
            continue
        idx, p = find_paragraph_by_prefix(doc, match)
        if idx is None:
            print(f"[{n}] MISS: prefix not found: {match[:60]!r}", file=sys.stderr)
            missed += 1
            continue
        old_text_preview = (p.text[:70] + "...") if len(p.text) > 70 else p.text
        if do_delete:
            print(f"[{n}] DELETE para {idx}: {old_text_preview!r}")
            if not dry_run:
                delete_paragraph(p)
        elif insert_after:
            if isinstance(insert_after, str):
                insert_after = [insert_after]
            print(f"[{n}] INSERT after para {idx} ({old_text_preview!r}):")
            for it in insert_after:
                pv = (it[:70] + "...") if len(it) > 70 else it
                print(f"         + {pv!r}")
            if not dry_run:
                insert_paragraphs_after(p, insert_after, smart_prop=smart_prop)
        else:
            new_text_preview = (replace[:70] + "...") if len(replace) > 70 else replace
            print(f"[{n}] para {idx}: {old_text_preview!r}\n         -> {new_text_preview!r}")
            if not dry_run:
                replace_paragraph_text(p, replace, smart_prop=smart_prop)
        applied += 1
    return applied, missed


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    g = ap.add_mutually_exclusive_group(required=True)
    g.add_argument("--apply", metavar="EDITS_JSON", help="apply edits")
    g.add_argument("--dry-run", metavar="EDITS_JSON", help="show what would change without writing")
    ap.add_argument("--smart-prop", action="store_true",
                    help="for proposition-formatted edits, rebuild bold-number+superscript-status formatting")
    ap.add_argument("--in", dest="docx_in", default=str(DOCX), help="input docx (default FORMLÆRE.docx)")
    ap.add_argument("--out", dest="docx_out", default=None, help="output docx (default: overwrite input)")
    args = ap.parse_args()

    edits_path = args.apply or args.dry_run
    is_dry = args.dry_run is not None
    edits_file = Path(edits_path)
    if not edits_file.exists():
        print(f"ERROR: edits file not found: {edits_file}", file=sys.stderr)
        return 2
    try:
        edits = json.loads(edits_file.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"ERROR: cannot parse edits JSON: {e}", file=sys.stderr)
        return 2
    if not isinstance(edits, list):
        print("ERROR: edits JSON must be a list of {match, replace} objects", file=sys.stderr)
        return 2

    doc_in = Path(args.docx_in)
    if not doc_in.exists():
        print(f"ERROR: docx not found: {doc_in}", file=sys.stderr)
        return 2
    doc = Document(str(doc_in))
    applied, missed = apply_edits(doc, edits, smart_prop=args.smart_prop, dry_run=is_dry)
    print(f"\n{'DRY-RUN' if is_dry else 'APPLY'}: {applied} applied, {missed} missed")
    if is_dry:
        return 0 if missed == 0 else 1

    out = Path(args.docx_out) if args.docx_out else doc_in
    doc.save(str(out))
    print(f"saved: {out}")
    return 0 if missed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
