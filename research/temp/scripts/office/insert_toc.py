#!/usr/bin/env python3
"""
Insert an auto-updating Word TOC field immediately before the Føreord
heading. The TOC pulls from Heading 1, 2 and 3.

When the resulting docx is opened in Word, right-click the placeholder
and choose 'Update field' to populate the TOC.

Run:
  python scripts/office/insert_toc.py
"""
from __future__ import annotations
import sys
from pathlib import Path

if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / 'FORMLÆRE.docx'


def make_para(text: str | None, style: str | None = None) -> 'OxmlElement':
    p_el = OxmlElement('w:p')
    if style:
        ppr = OxmlElement('w:pPr')
        ps = OxmlElement('w:pStyle')
        ps.set(qn('w:val'), style)
        ppr.append(ps)
        p_el.append(ppr)
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        p_el.append(r)
    return p_el


def build_toc_field_paragraph() -> 'OxmlElement':
    p = OxmlElement('w:p')

    # field begin
    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1)
    p.append(r1)

    # instruction text — use raw string to avoid backslash escapes
    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '
    r2.append(instr)
    p.append(r2)

    # field separate
    r3 = OxmlElement('w:r')
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r3.append(fc2)
    p.append(r3)

    # placeholder visible content
    r4 = OxmlElement('w:r')
    t4 = OxmlElement('w:t')
    t4.text = '(Innhaldslista vert oppdatert i Word: høgreklikk her og vel «Oppdater felt».)'
    r4.append(t4)
    p.append(r4)

    # field end
    r5 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r5.append(fc3)
    p.append(r5)

    return p


def build_page_break() -> 'OxmlElement':
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def main() -> int:
    d = Document(str(DOCX))
    foreord_idx = None
    foreord_p = None
    for i, p in enumerate(d.paragraphs):
        if p.text.strip() == 'Føreord':
            foreord_idx = i
            foreord_p = p
            break
    if foreord_p is None:
        print('ERROR: Føreord paragraph not found', file=sys.stderr)
        return 1

    # Check if TOC already exists (idempotent)
    for p in d.paragraphs[:foreord_idx]:
        if p.text.strip() == 'Innhald':
            print('Innhald heading already present before Føreord — skipping insertion')
            return 0

    foreord_el = foreord_p._p
    parent = foreord_el.getparent()
    pos = list(parent).index(foreord_el)

    innhald_heading = make_para('Innhald', style='Heading1')
    toc_field = build_toc_field_paragraph()
    page_break = build_page_break()

    parent.insert(pos, innhald_heading)
    parent.insert(pos + 1, toc_field)
    parent.insert(pos + 2, page_break)

    d.save(str(DOCX))
    print(f'Inserted Innhald heading + TOC field + page break before Føreord (was at index {foreord_idx})')
    return 0


if __name__ == '__main__':
    sys.exit(main())
