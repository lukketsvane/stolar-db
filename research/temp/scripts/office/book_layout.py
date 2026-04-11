#!/usr/bin/env python3
"""
Apply Wittgenstein-Tractatus book layout to FORMLÆRE.docx.

Models the layout from the Norwegian Tractatus (Gyldendal 1999, digital
facsimile in digibok_2009032304095/):

  - Small page format (~12 × 19 cm)
  - Generous margins
  - Hanging-indent propositions: number in left column at 0, body text
    starting at fixed column ~28 mm. ALL depths use the SAME indent. The
    decimal number itself shows the depth.
  - Chapter headings in SMALL CAPS, centred, with significant vertical
    space above/below
  - Centred page numbers in footer
  - Tab character between proposition number+status and body text

Run:
  python scripts/office/book_layout.py
  python scripts/office/book_layout.py --dry-run
"""
from __future__ import annotations
import argparse
import re
import sys
from pathlib import Path

if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

from docx import Document
from docx.shared import Cm, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / 'FORMLÆRE_latest.docx'

# Page format — Tractatus translation (Gyldendal 1999, digibok measurement)
# Pages are 1952 × 3200 px ≈ 124 × 203 mm at 400 dpi. Aspect ratio 1.64.
# Round to a clean Norwegian small-book format.
PAGE_W_MM = 125
PAGE_H_MM = 200
MARGIN_TOP_MM = 16
MARGIN_BOTTOM_MM = 18
MARGIN_LEFT_MM = 18
MARGIN_RIGHT_MM = 18

# Hanging-indent geometry for propositions
# Longest number is "5.521" (5 chars) + status letter as superscript.
# 22 mm comfortably accommodates that at 11pt Garamond and leaves a small gap
# before the body text begins.
NUMBER_COL_MM = 22
BODY_COL_OFFSET_MM = NUMBER_COL_MM  # body text starts here

# Patterns
PROP_RE = re.compile(r'^(\d+(?:\.\d+)*)\s*([daiot])\s+(.*)$', re.DOTALL)
APPENDIX_PROP_RE = re.compile(r'^([DTA])(\d+)\s+(.*)$', re.DOTALL)


def set_page_format(doc):
    sec = doc.sections[0]
    sec.page_width = Mm(PAGE_W_MM)
    sec.page_height = Mm(PAGE_H_MM)
    sec.top_margin = Mm(MARGIN_TOP_MM)
    sec.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    sec.left_margin = Mm(MARGIN_LEFT_MM)
    sec.right_margin = Mm(MARGIN_RIGHT_MM)


def apply_small_caps(run):
    """Apply small caps to a run via direct XML."""
    rpr = run._element.get_or_add_rPr()
    sc = rpr.find(qn('w:smallCaps'))
    if sc is None:
        sc = OxmlElement('w:smallCaps')
        rpr.append(sc)
    sc.set(qn('w:val'), '1')


def set_letter_spacing(run, twips: int):
    """Apply character spacing in twips (1/20 pt)."""
    rpr = run._element.get_or_add_rPr()
    sp = rpr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        rpr.append(sp)
    sp.set(qn('w:val'), str(twips))


def style_heading(p, *, level: str):
    """Style a heading paragraph as centred small caps with vertical spacing.
    level is 'h1' (chapter), 'h2' (section), 'h3' (subsection)."""
    sizes = {'h1': 13, 'h2': 11, 'h3': 10}
    space_before = {'h1': 24, 'h2': 16, 'h3': 10}
    space_after = {'h1': 14, 'h2': 8, 'h3': 6}
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = None
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(space_before[level])
    p.paragraph_format.space_after = Pt(space_after[level])
    for r in p.runs:
        r.font.name = 'Garamond'
        r.font.size = Pt(sizes[level])
        r.bold = False
        r.italic = False
        apply_small_caps(r)
        set_letter_spacing(r, 30)  # ~1.5pt extra spacing for elegant small caps


def style_proposition(p, *, number_col_mm: int = NUMBER_COL_MM):
    """Apply hanging indent + tab stop and replace the space between
    'N.N x' and the body with a tab character. All depths get the SAME
    indent — the number itself encodes the depth."""
    text = p.text
    m = PROP_RE.match(text)
    if not m:
        return False
    num, status, body = m.groups()
    # Set hanging indent on paragraph
    pf = p.paragraph_format
    pf.left_indent = Mm(number_col_mm)
    pf.first_line_indent = -Mm(number_col_mm)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    # Add a tab stop at the body column
    tab_stops = pf.tab_stops
    # Clear existing tab stops and add a fresh one at the body column
    tab_stops.clear_all()
    tab_stops.add_tab_stop(Mm(number_col_mm), WD_TAB_ALIGNMENT.LEFT)
    # Rebuild runs: number bold, status superscript, tab, body
    para_el = p._p
    for child in list(para_el):
        if child.tag == qn('w:r'):
            para_el.remove(child)
    # Run 1: number (bold)
    r1 = p.add_run(num)
    r1.bold = True
    r1.font.name = 'Garamond'
    r1.font.size = Pt(11)
    # Run 2: status letter (small superscript)
    r2 = p.add_run(status)
    r2.bold = True
    r2.font.name = 'Garamond'
    r2.font.size = Pt(11)
    r2.font.superscript = True
    # Run 3: tab + body text
    r3 = p.add_run('\t' + body)
    r3.font.name = 'Garamond'
    r3.font.size = Pt(11)
    return True


def style_appendix_block(p, *, number_col_mm: int = NUMBER_COL_MM):
    """Style appendix entries (D4, T6, A1) with the same hanging indent
    so they line up with the propositions."""
    text = p.text
    m = APPENDIX_PROP_RE.match(text)
    if not m:
        return False
    letter, num, body = m.groups()
    label = f'{letter}{num}'
    pf = p.paragraph_format
    pf.left_indent = Mm(number_col_mm)
    pf.first_line_indent = -Mm(number_col_mm)
    pf.space_before = Pt(8)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(Mm(number_col_mm), WD_TAB_ALIGNMENT.LEFT)
    para_el = p._p
    for child in list(para_el):
        if child.tag == qn('w:r'):
            para_el.remove(child)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = 'Garamond'
    r1.font.size = Pt(11)
    r2 = p.add_run('\t' + body)
    r2.font.name = 'Garamond'
    r2.font.size = Pt(11)
    return True


def style_body_paragraph(p):
    """Set body font + size on a normal paragraph (not a heading or
    proposition). Used for prose paragraphs in føreord, etterord, etc.
    Resets any progressive indent left over from the old indent.py."""
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = None
    pf.line_spacing = 1.15
    if pf.space_before is None:
        pf.space_before = Pt(4)
    if pf.space_after is None:
        pf.space_after = Pt(4)
    for r in p.runs:
        if r.font.name is None or 'Calibri' in (r.font.name or ''):
            r.font.name = 'Garamond'
        if r.font.size is None:
            r.font.size = Pt(11)


CHAPTER_TITLES = {
    '1 Form er ein posisjon i eit rom av moglegheiter.': 'h1',
    '2 Ikkje alle posisjonar er like sannsynlege.': 'h1',
    '3 Seleksjonstrykka produserer eit landskap over formrommet.': 'h1',
    '4 Landskapet er dynamisk.': 'h1',
    '5 Det finst agentar som responderer på landskapet.': 'h1',
    '6 Forma oppstår mellom agentane.': 'h1',
    '7 Ingen form er endeleg.': 'h1',
    '7 Ingen form er endeleg; navigasjonen held fram.': 'h1',
}

SECTION_NAMES = {
    'Innhald': 'h1',
    'Føreord': 'h1',
    'Ordliste': 'h1',
    'Etterord': 'h1',
    'Referansar': 'h1',
}


def is_heading(p) -> str | None:
    """Return level ('h1', 'h2', 'h3') or None if p is not a heading.

    Conservative classifier: trusts the existing docx style names first,
    and only falls back to a hardcoded whitelist of known chapter and
    section titles. This avoids misclassifying body paragraphs that
    happen to start with a number."""
    s = p.style.name if p.style else ''
    t = p.text.strip()
    if not t:
        return None

    # Word style heading takes priority
    if s.startswith('Heading 1'):
        return 'h1'
    if s.startswith('Heading 2'):
        return 'h2'
    if s.startswith('Heading 3'):
        # In FORMLÆRE_latest the chapter headings (1-7) are wrongly tagged
        # as Heading 3. Promote them to h1 if the text matches a known title.
        if t in CHAPTER_TITLES:
            return 'h1'
        return 'h2'
    if s.startswith('Heading 4'):
        if t in CHAPTER_TITLES:
            return 'h1'
        # A.x sections sometimes get tagged as Heading 4
        if re.match(r'^A\.\d\s', t):
            return 'h2'
        return 'h2'
    if s.startswith('Heading 5'):
        # A.6.x sub-sections
        if re.match(r'^A\.\d\.\d\s', t):
            return 'h3'
        return 'h3'

    # Style is normal — only classify if text is a known section/chapter title
    if t in SECTION_NAMES:
        return SECTION_NAMES[t]
    if t in CHAPTER_TITLES:
        return CHAPTER_TITLES[t]
    if t.startswith('A  Formell') or t.startswith('A Formell'):
        return 'h1'
    if re.match(r'^A\.\d\.\d\s', t):
        return 'h3'
    if re.match(r'^A\.\d\s', t):
        return 'h2'
    return None


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument('--dry-run', action='store_true')
    args = ap.parse_args()

    d = Document(str(DOCX))
    set_page_format(d)

    n_props = 0
    n_appendix = 0
    n_headings = 0
    n_body = 0

    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if not t:
            continue

        level = is_heading(p)
        if level is not None:
            if not args.dry_run:
                style_heading(p, level=level)
            n_headings += 1
            print(f'  {level} para {i:3d}: {t[:70]}')
            continue

        # Try proposition
        if PROP_RE.match(t):
            if not args.dry_run:
                style_proposition(p)
            n_props += 1
            continue

        # Try formal appendix entry (D, T, A)
        if APPENDIX_PROP_RE.match(t):
            if not args.dry_run:
                style_appendix_block(p)
            n_appendix += 1
            continue

        # Body / prose paragraph: clear progressive indents
        if not args.dry_run:
            style_body_paragraph(p)
        n_body += 1

    if not args.dry_run:
        d.save(str(DOCX))

    print(f"\n{'DRY' if args.dry_run else 'APPLY'}: "
          f"{n_headings} headings, {n_props} propositions, "
          f"{n_appendix} appendix entries, {n_body} body paragraphs")
    return 0


if __name__ == '__main__':
    sys.exit(main())
