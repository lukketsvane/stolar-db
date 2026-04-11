#!/usr/bin/env python3
"""
Rebuild A.6 Empiriske testresultat as a clean catalog of strong findings,
each with one figure embedded inline.
"""
from __future__ import annotations
import sys
from pathlib import Path
from copy import deepcopy

if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / 'FORMLÆRE.docx'
FIG = ROOT / 'analysis' / 'figures'

# (subsection heading, body text, figure file, italic caption)
SECTIONS = [
    (
        'A.6.1  Morphospace ikkje uniformt (1.4)',
        'Nærmaste-nabo-distansen i (H, W, D) etter z-skalering har CV = 5,4 (95 % CI [3,7, 5,9]), om lag 15× over Poisson-nullen 0,36. Funna held i alle 14 hold-out-subset (museum, periode, stil), inkludert NMK aleine (n = 63). n = 1664.',
        'fig-1.4-nn-cv.png',
        'Nærmaste-nabo-fordelinga er klumpa, ikkje uniform.',
    ),
    (
        'A.6.2  Stilperiode som samlevariabel (2.4, 2.62)',
        'Gjensidig informasjon (sklearn k-NN) mellom prediktor og kvar geometrisk dimensjon: stilperiode slår grov materialgruppe på alle fire (H, W, D, H/W). Gevinsten er stor: stil 0,30–0,59 bits mot mat 0,06–0,10 bits. Forholdstal opp til 7×. Resultatet held under museum-CV og under leave-one-period-out. n = 1664.',
        'fig-2.4-proxy.png',
        'Stilperiode slår materiale på alle fire dimensjonar.',
    ),
    (
        'A.6.3  Kanaliseringshierarki i mesh-rommet (3.3)',
        'Variasjonskoeffisienten på tvers av seks mesh-trekk strekkjer seg over to storleiksordnar. Sphericity er det mest kanaliserte trekkjet (CV = 0,074); råvolumet det friaste (CV = 9,5). Spreiinga er 128× (95 % CI [50×, 158×]). n = 2202.',
        'fig-3.3-channeling.png',
        'Kanaliseringshierarkiet i mesh-rommet: sphericity er sterkt kanalisert, råvolumet fritt, ein 128× spreiing.',
    ),
    (
        'A.6.4  Stilar er gradientar, ikkje topologiske klynger (3.4)',
        'Silhouette-skoren for stilperiode i 4D mesh-trekk-rom (sphericity, fill_ratio, inertia_ratio, complexity) er −0,34 (95 % CI [−0,37, −0,33]) over 25 stilar med minst 10 medlemar kvar. Negativ silhouette betyr at gjennomsnittspunktet i ein stil ligg nærmare punkta i naboklynga enn dei eigne. Stilkategoriane er gradientar, ikkje topologisk skilde regionar. n = 1971.',
        'fig-3.4-silhouette.png',
        'Stilar er gradientar i mesh-rommet, ikkje topologiske klynger.',
    ),
    (
        'A.6.5  Kumulativ ekspansjon av formrommet (4.4)',
        'Det kumulative konvekse hylsterveolumet i (H, W, D), etter klipping til 1.–99. persentil per dimensjon, veks monotont gjennom 24 femtjueårs-periodar. Totalvekst 107× (95 % CI [30×, breitt]). I mesh-trekk-rommet er den same testen 553×. Landskapet skrumpar aldri.',
        'fig-4.4-hull.png',
        'Det kumulative formromsvolumet veks monotont gjennom 700 år.',
    ),
    (
        'A.6.6  Mahogni-kollapsen 1825-1849 (4.5)',
        'Av norskproduserte stolar i perioden 1825-1849 inneheld 16 av 16 mahogni (deterministisk). I perioden 1750-1799 var fraksjonen null. H/W-variasjonskoeffisienten i kollaps-perioden er 0,083, mot 0,140 i 1750-1799 og 0,090 i 1850-1899. Eitt seleksjonstrykk vart så dominant at både materialrommet og formrommet kollapsa.',
        'fig-4.5-mahogni.png',
        'Norsk mahogni-kollapsen 1825-1849: eitt seleksjonstrykk overkøyrer alle andre.',
    ),
    (
        'A.6.7  Direkte falsifisering av postulat 4.1',
        'Wasserstein-distansen mellom suksessive 50-årsperiodar gjev mean 14,4 cm for høgde, 8,2 cm for breidde og 5,9 cm for djupn. Ingen av dei ti periodepara har distanse under 0,5 cm. Postulatet om at landskapet endrar seg held mot direkte falsifisering. Same metode på random-walk-null forkastar denne med p ≪ 10⁻⁶³ for kvar dimensjon.',
        'fig-falsification-4.1.png',
        'Wasserstein-distanse mellom suksessive 50-årsperiodar: landskapet endrar seg overalt.',
    ),
]


def make_paragraph_with_text(template_p, text: str, *, style: str | None = None) -> 'OxmlElement':
    """Build a new w:p element copied from template, with one run of text."""
    el = deepcopy(template_p._p)
    for child in list(el):
        if child.tag == qn('w:r') or child.tag == qn('w:fldSimple'):
            el.remove(child)
    if style:
        # set pStyle in pPr
        ppr = el.find(qn('w:pPr'))
        if ppr is None:
            ppr = OxmlElement('w:pPr')
            el.insert(0, ppr)
        existing = ppr.find(qn('w:pStyle'))
        if existing is not None:
            ppr.remove(existing)
        ps = OxmlElement('w:pStyle')
        ps.set(qn('w:val'), style)
        ppr.insert(0, ps)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    el.append(r)
    return el


def main() -> int:
    d = Document(str(DOCX))

    # Locate A.6 heading and Referansar (or end of doc)
    start = end = None
    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if t.startswith('A.6') and 'Empiriske' in t:
            start = i
        elif start is not None and t == 'Referansar':
            end = i
            break
    if start is None:
        print('ERROR: A.6 heading not found', file=sys.stderr); return 1
    if end is None:
        end = len(d.paragraphs)
    print(f'A.6 region: {start}..{end - 1}')

    # Use a body paragraph from the doc as template (e.g. para 12 from the føreord)
    template_p = d.paragraphs[12] if len(d.paragraphs) > 12 else d.paragraphs[start + 1]

    # Capture A.6 heading element (we'll keep this) and delete everything between it and Referansar
    a6_heading_p = d.paragraphs[start]
    parent = a6_heading_p._p.getparent()
    paras = list(d.paragraphs)
    to_delete = [paras[i]._p for i in range(start + 1, end)]
    print(f'Deleting {len(to_delete)} old A.6 body paragraphs')
    for el in to_delete:
        el.getparent().remove(el)

    # Now insert new A.6 content right after the A.6 heading
    # Insert intro paragraph + each subsection (heading + body + image + caption)
    insert_after = a6_heading_p._p
    insert_pos = list(parent).index(insert_after) + 1

    intro_text = (
        'Kvar testresultat under er ein direkte måling mot ein nullhypotese. '
        'Berre funn med 95 % bootstrap-konfidensintervall som ikkje krysser '
        'nullen, og som held i alle hold-out-subset, er inkluderte. Alle '
        'figurar har éi-linje overskrift; tala finst i tabellforma i '
        'analysis/evidence_table.csv.'
    )
    elements_to_insert: list = []
    elements_to_insert.append(make_paragraph_with_text(template_p, intro_text))

    for heading_text, body_text, fig_name, caption_text in SECTIONS:
        # Heading (will be classified as h3 by headings.py — A.6.x pattern)
        elements_to_insert.append(make_paragraph_with_text(template_p, heading_text))
        # Body
        elements_to_insert.append(make_paragraph_with_text(template_p, body_text))
        # Image placeholder (we'll fill the picture below)
        img_p = deepcopy(template_p._p)
        for child in list(img_p):
            if child.tag == qn('w:r') or child.tag == qn('w:fldSimple'):
                img_p.remove(child)
        elements_to_insert.append(img_p)
        # Caption
        cap_p = make_paragraph_with_text(template_p, caption_text)
        elements_to_insert.append(cap_p)

    # Insert all in reverse so insert_pos remains the anchor point
    for el in reversed(elements_to_insert):
        parent.insert(insert_pos, el)

    d.save(str(DOCX))

    # Reopen so we can use python-docx wrappers to add pictures and style captions
    d = Document(str(DOCX))
    # Find the rebuilt A.6 region
    new_start = None
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith('A.6') and 'Empiriske' in p.text:
            new_start = i
            break
    # Walk forward through the inserted blocks
    cursor = new_start + 1  # skip A.6 heading
    cursor += 1  # skip intro paragraph
    for heading_text, body_text, fig_name, caption_text in SECTIONS:
        # heading at cursor
        heading_p = d.paragraphs[cursor]; cursor += 1
        # body at cursor
        body_p = d.paragraphs[cursor]; cursor += 1
        # image placeholder at cursor — fill it
        img_para = d.paragraphs[cursor]; cursor += 1
        img_path = FIG / fig_name
        if img_path.exists():
            img_run = img_para.add_run()
            img_run.add_picture(str(img_path), width=Cm(14))
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.left_indent = None
            img_para.paragraph_format.space_before = Pt(8)
            img_para.paragraph_format.space_after = Pt(2)
        # caption at cursor — italicize
        cap_para = d.paragraphs[cursor]; cursor += 1
        for r in cap_para.runs:
            r.italic = True
            r.font.size = Pt(9)
            r.font.name = 'Garamond'
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.left_indent = None
        cap_para.paragraph_format.space_before = Pt(0)
        cap_para.paragraph_format.space_after = Pt(10)

    d.save(str(DOCX))
    print('A.6 rebuilt with 7 figured sub-sections')
    return 0


if __name__ == '__main__':
    sys.exit(main())
