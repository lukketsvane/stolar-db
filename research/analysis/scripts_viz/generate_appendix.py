import pandas as pd
import numpy as np
import docx
from docx.shared import Inches, Pt
import os
import warnings
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from scipy.stats import gaussian_kde
from scipy.ndimage import maximum_filter
from docx.oxml.ns import qn
from viz_style import setup_style, finalize_plot

warnings.filterwarnings('ignore')
setup_style()

# Regenerate multimodal plot
df = pd.read_csv('STOLAR/STOLAR.csv')
df = df.dropna(subset=['Høgde (cm)', 'Breidde (cm)'])
df = df[df['Breidde (cm)'] > 0]
df = df[df['Høgde (cm)'] > 0]

fig = plt.figure(figsize=(10, 8))
ax = fig.add_subplot(111, projection='3d')
ax.set_facecolor('#ffffff')
x, y = df['Breidde (cm)'], df['Høgde (cm)']
xmin, xmax, ymin, ymax = 20, 140, 40, 170
X, Y = np.mgrid[xmin:xmax:100j, ymin:ymax:100j]
positions = np.vstack([X.ravel(), Y.ravel()])
values = np.vstack([x, y])
kernel = gaussian_kde(values)
Z = np.reshape(kernel(positions).T, X.shape)
Z = Z / Z.max()
surf = ax.plot_surface(X, Y, Z, cmap='Greys', edgecolor='#cccccc', linewidth=0.1, antialiased=True, alpha=0.8)
data_max = maximum_filter(Z, 10)
maxima = (Z == data_max) & (Z > 0.1)
max_idx = np.where(maxima)
ax.scatter(X[max_idx], Y[max_idx], Z[max_idx], color='black', s=50, zorder=10, marker='*')
ax.set_xlabel('Breidde (cm)', fontname='EB Garamond')
ax.set_ylabel('Høgde (cm)', fontname='EB Garamond')
ax.set_zlabel('Tettleik', fontname='EB Garamond')
ax.view_init(elev=35, azim=-45)
ax.xaxis.pane.fill = ax.yaxis.pane.fill = ax.zaxis.pane.fill = False
plt.tight_layout()
plt.savefig('analysis/figures_new/fig_3_2_multimodal.png', dpi=300)
plt.close()

# --- DOCX Generation ---
doc = docx.Document()

def set_font(run, size=11, bold=False):
    run.font.name = 'EB Garamond'
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element.get_or_add_rPr()
    rFonts = r.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'EB Garamond')
    rFonts.set(qn('w:hAnsi'), 'EB Garamond')

def add_styled_heading(doc, text, level):
    h = doc.add_heading('', level)
    run = h.add_run(text)
    set_font(run, size=18 if level==1 else 14, bold=True)

doc.add_heading('FORMLÆRE: Empirisk Appendiks', 0)

# Section A.6: Empiriske testresultat
add_styled_heading(doc, 'A.6 Empiriske testresultat', 1)

sections = [
    ("A.6.1 Formrommet er ikkje uniformt busett (1.4)", "Nærmaste-nabo-distansen i (H, W, D) etter z-skalering har ein variasjonskoeffisient (CV) på 5,4 (95 % CI [3,7, 5,9]), noko som er om lag 15 gonger over Poisson-nullhypotesen på 0,36. Funna er robuste på tvers av 14 hold-out-subset (museum, periode, stil), inkludert NMK isolert (n = 63). n = 1664.", 'analysis/figures_new/fig_1_4_uniformitet.png'),
    ("A.6.2 Stilperiode som samlevariabel (2.4, 2.62)", "Gjensidig informasjon (sklearn k-NN) mellom prediktor og kvar geometrisk dimensjon: stilperiode slår grov materialgruppe på alle fire (H, W, D, H/W). Gevinsten er stor: stil 0,30–0,59 bits mot mat 0,06–0,10 bits. Forholdstal opp til 7×. Resultatet held under museum-CV og under leave-one-period-out. n = 1664.", 'analysis/figures_new/fig_2_4_mi_style_mat.png'),
    ("A.6.3 Kanaliseringshierarki i mesh-rommet (3.3)", "Variasjonskoeffisienten på tvers av seks mesh-trekk strekkjer seg over to storleiksordnar. Sphericity er det mest kanaliserte trekkjet (CV = 0,074); råvolumet det friaste (CV = 9,5). Spreiinga er 128× (95 % CI [50×, 158×]). n = 2202.", 'analysis/figures_new/mesh_3_3_channeling.png'),
    ("A.6.4 Stilar er gradientar, ikkje topologiske klynger (3.4)", "Silhouette-skoren for stilperiode i 4D mesh-trekk-rom (sphericity, fill_ratio, inertia_ratio, complexity) er −0,34 (95 % CI [−0,37, −0,33]) over 25 stilar med minst 10 medlemar kvar. Negativ silhouette betyr at gjennomsnittspunktet i ein stil ligger nærmare punkta i naboklynga enn dei eigne. Stilkategoriane er gradientar, ikkje topologisk skilde regionar. n = 1971.", 'analysis/figures_new/mesh_3_4_silhouette.png'),
    ("A.6.5 Kumulativ ekspansjon av formrommet (4.4)", "Det kumulative konvekse hylsterveolumet i (H, W, D), etter klipping til 1.–99. persentil per dimensjon, veks monotont gjennom 24 femtiårs-periodar. Totalvekst 107× (95 % CI [30×, breitt]). I mesh-trekk-rommet er den same testen 553×. Landskapet skrumpar aldri.", 'analysis/figures_new/mesh_4_4_hull.png'),
    ("A.6.6 Mahogni-kollaps 1825-1849 (4.5)", "iI utvalet av norskproduserte stolar frå perioden 1825–1849 er bruken av mahogni deterministisk (16 av 16), samanlikna med ein fraksjon på null i perioden 1750–1799. Variasjonskoeffisienten for H/W i kollaps-perioden er 0,083, mot 0,140 i den føregåande perioden. Dette illustrerer korleis eitt seleksjonstrykk kan verte så dominant at formrommet kollapsar.", None),
    ("A.6.7 Direkte falsifisering av postulat 4.1", "Wasserstein-distansen mellom suksessive 50-årsperiodar gjev mean 14,4 cm for høgde, 8,2 cm for breidde og 5,9 cm for djupn. Ingen av dei ti periodepara har distanse under 0,5 cm. Postulatet om at landskapet endrar seg held mot direkte falsifisering. Same metode på random-walk-null forkastar denne med p ≪ 10⁻⁶³ for kvar dimensjon.", None),
    ("A.6.8 Tilpassingsfunksjonen har fleire haugar (3.2)", "Kjernedestitetsestimering (KDE) i todimensjonalt formrom (H, Breidde) syner fleire distinkte lokale maksimum (haugar) skild av dalar. Landskapet er multimodalt, ikkje ein unimodal normalfordeling.", 'analysis/figures_new/fig_3_2_multimodal.png'),
    ("A.6.9 Endringa i formrommet er diskontinuerleg (4.3)", "Distribusjonen av endringsratar (hopp i median mellom 50-årsperiodar) viser at stase vert avbrote av brå topologiske skifte. Maksimum-hoppet er 6,08 gonger høgare enn median-endringa.", 'analysis/figures_new/fig_4_3_stase_brot.png'),
    ("A.6.10 Formgjeving er ikkje ei tilfeldig vandring (5.1)", "Kolmogorov-Smirnov-test for (H, W, D) forkastar nullhypotesen om uniform fordeling (p ≪ 0,001). Variasjonen er underlagt styring mot attraktorar, drive av negativ tilbakekopling.", 'analysis/figures_new/fig_5_1_uniform.png'),
    ("A.6.13 3D Shape Grammar som spesialtilfelle av FORMLÆRE", "Metodikken frå Xue og Chen (2024) er bygd ut og implementert direkte på 3D-mesh-samlinga (N = 2000). Eit genbasseng av 8 medoid-stolar dannar grunnlaget for generering av teoretiske mutasjonar via erstatning, skalering og forskyving. Resultatet syner at dei teoretisk genererte formene utforskar det topologiske moglegheitsrommet. Dette er ein empirisk stadfesting av proposisjon 3.42: Shape grammar-algoritmar opererer innanfor det formaliserte rammeverket som eit spesialtilfelle.", 'analysis/figures_new/fig_shape_grammar_kde.png'),
    ("A.6.14 Tilpassingslandskap og Mean Reversion (3.1)", "Ein test av Ornstein-Uhlenbeck-prosessen plottar avvik frå det globale sentrumet (medianhøgde) mot endringa i den komande 25-årsperioden. Ein klår negativ korrelasjon (helling på -0,32) syner tilbakevending mot gjennomsnittet (mean reversion). Former som driv langt vekk frå attraktorane kjenner eit seleksjonstrykk som trekkjer dei tilbake mot kjernen over tid.", 'analysis/figures_new/fig_3_1_ou_reversion.png'),
    ("A.6.15 Vektorfelt og stiavhengigheit (4.3, 4.5)", "Aggregering av dei morfologiske forflyttingane i rommet (breidde og djupn) frå éin 50-årsperiode til den neste dannar eit tydeleg vektorfelt. Stolane spreier seg ikkje isotropisk, men fylgjer distinkte straumlinjer mot lokale optima. Dette er empirisk prov for stiavhengigheit (Arthur 1994).", 'analysis/figures_new/fig_path_dependence_flow.png'),
    ("A.6.16 Affin-invariant disparitetsvekst (4.4)", "Den kumulative ekspansjonen av formrommet kan òg påvisast uavhengig av euklidske metrikkar ved bruk av generalisert varians. Dette målet er invariant under affine transformasjonar.", 'analysis/figures_new/fig_affine_disparity.png'),
    ("A.6.17 Agensi og negativ tilbakekopling (5.1)", "Måling av lag-1 autokorrelasjon i formavvik innanfor einskilde stiltradisjonar over tid gjev ein negativ verdi (r = -0.42). Den negative autokorrelasjonen provar homeostatisk feilkorrigering.", 'analysis/figures_new/fig_levin_feedback.png'),
    ("A.6.18 Nisjepartisjonering (5.3)", "Måling av konvekse hylster for ulike funksjonelle klassar (som lenestol, krakk og standard stol) syner at dei utfyller kvarandre i rommet med minimal overlapping. Slik nisjepartisjonering bekreftar at utforminga i seg sjølv konstruerer nisjar (Odling-Smee et al. 2003).", 'analysis/figures_new/fig_niche_overlap.png'),
    ("A.6.19 Det tilstøytande moglege og kompleksitetstrakta (6.5)", "Analyse av maskevidde (mesh complexity) frå 1600 til 1950 syner ei kompleksitetstrakt som gradvis utvidar seg over tid. Dette bekreftar teorien om at nyskaping skjer ved å utforske «det tilstøytande moglege» (Kauffman 1993).", 'analysis/figures_new/fig_complexity_funnel.png'),
    ("A.6.20 Lokal varians og forsterkande seleksjon (7.2)", "Relasjonen mellom lokal punkttettleik og lokal morfologisk varians syner ein 'varians-felle': høgare tettleik (suksess) gjev lågare lokal varians (konformitet). Klynger fungerer som sterke attraktorar der nyskaping minkar. Visualiseringa samanliknar grid av stol-rektangel frå høg-tettleiks (venstre) vs låg-tettleiks regionar (høgre).", 'analysis/figures_new/fig_density_variance.png')
]

for title, text, img in sections:
    add_styled_heading(doc, title, 2)
    tab = doc.add_table(rows=1, cols=2)
    tab.autofit = False
    tab.columns[0].width = Inches(4.2)
    tab.columns[1].width = Inches(2.3)
    set_font(tab.rows[0].cells[0].paragraphs[0].add_run(text))
    if img and os.path.exists(img):
        run = tab.rows[0].cells[1].paragraphs[0].add_run()
        run.add_picture(img, width=Inches(2.2))

doc.save('empiri_appendiks.docx')
